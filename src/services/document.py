"""Document generation, composition, and file content extraction service."""

import io
import os
import pathlib
import re
from typing import Any, Dict, Optional

from src.core.transparency_logger import transparency_logger
from src.utils.logger import get_logger

logger = get_logger(__name__)


def _parse_outline_sections(outline: str) -> list:
    """Parse numbered sections from an outline into a list of {number, title, bullets} dicts."""
    sections = []
    current: Optional[Dict[str, Any]] = None
    for line in outline.split("\n"):
        match = re.match(r"^\s*(\d+)[\.\)]\s+(.+)", line)
        if match:
            if current:
                sections.append(current)
            current = {
                "number": int(match.group(1)),
                "title": match.group(2).strip(),
                "bullets": [],
            }
        elif current and re.match(r"^\s*[-•*]\s+.+", line):
            current["bullets"].append(line.strip())
    if current:
        sections.append(current)
    return sections


async def compose_document(
    topic: str,
    instructions: str,
    context: str = "",
    format: str = "markdown",
    length: str = "short",
    settings_service=None,
    conversation_id: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Compose a long-form document using a multi-step LLM pipeline.

    For 'short': plan → write (single pass) → review  (3 LLM calls)
    For 'medium': plan → write each section → assemble → review  (2 + N + 1 LLM calls)
    For 'long': plan (detailed) → write each section with depth → assemble → review  (2 + N + 1 LLM calls)

    Args:
        topic: What to write about
        instructions: Tone, audience, length, format, key points
        context: Research findings, reference material, background
        format: Output format ("markdown" or "plain")
        length: Output length — "short" (default), "medium", or "long"
        settings_service: SettingsService instance for LLM configuration

    Returns:
        Dict with keys: content (str), outline (str), sections (int)
    """
    if not settings_service:
        raise ValueError("settings_service is required for compose_document")

    from src.core.llm_client import LLMClient, LLMConfig

    # Build LLM client from settings, using the writer model if configured
    llm_config = LLMConfig.from_settings(settings_service)
    writer_model = llm_config.get_writer_model()
    llm_client = LLMClient(llm_config)

    format_instruction = (
        "Use markdown formatting (headings, lists, bold, etc.)."
        if format == "markdown"
        else "Use plain text without any markdown formatting."
    )

    context_block = f"\n\nReference material / context:\n{context}" if context else ""

    # --- Length configuration ---
    if length == "long":
        section_guidance = (
            "Create 8-12 sections. For each section include 3-5 detailed bullet points "
            "covering key concepts, examples, and sub-topics."
        )
        section_tokens = 6144
        depth_instruction = (
            " Write in depth with detailed explanations, concrete examples, and thorough coverage. "
            "Each section should be comprehensive and self-contained."
        )
    elif length == "medium":
        section_guidance = (
            "Create 6-8 sections. For each section include 3-4 bullet points covering key concepts."
        )
        section_tokens = 4096
        depth_instruction = " Write with good detail and clear explanations."
    else:  # short
        section_guidance = "Include section titles and 2-3 key bullet points for each section."
        section_tokens = None  # not used in single-pass mode
        depth_instruction = ""

    # --- Step 1: Plan ---
    logger.info(f"compose_document [{length}]: planning outline for '{topic[:100]}'")
    plan_prompt = (
        f"Topic: {topic}\n"
        f"Instructions: {instructions}\n"
        f"{context_block}\n\n"
        f"Produce a structured outline for this document. "
        f"{section_guidance} "
        "Number each section."
    )
    outline = llm_client.complete_text(
        prompt=plan_prompt,
        system_message="You are a document planning specialist. Create clear, well-organized outlines.",
        temperature=0.5,
        max_tokens=2048,
        model_override=writer_model,
    )
    logger.info("compose_document: outline complete")

    # Persist the outline as an internal transparency row (visibility only;
    # billing reads llm_consumption). Billing-neutral.
    if conversation_id:
        transparency_logger.log(
            conversation_id, "document", outline, role="assistant",
            extra_metadata={"stage": "outline"},
        )

    if length == "short":
        # --- Step 2 (short): Write full document in a single pass ---
        logger.info("compose_document [short]: writing full document")
        write_prompt = (
            f"Topic: {topic}\n"
            f"Instructions: {instructions}\n"
            f"{context_block}\n\n"
            f"Outline to follow:\n{outline}\n\n"
            f"Write the full document following the outline above. "
            f"Complete ALL sections thoroughly. {format_instruction}"
        )
        draft = llm_client.complete_text(
            prompt=write_prompt,
            system_message="You are an expert writer. Produce well-structured, comprehensive content.",
            temperature=0.7,
            max_tokens=8192,
            model_override=writer_model,
        )
        logger.info("compose_document [short]: draft complete")

    else:
        # --- Step 2 (medium/long): Write each section individually ---
        sections = _parse_outline_sections(outline)
        if not sections:
            # Fallback: treat as single-pass if outline couldn't be parsed
            logger.warning(
                "compose_document: could not parse sections, falling back to single-pass"
            )
            write_prompt = (
                f"Topic: {topic}\n"
                f"Instructions: {instructions}\n"
                f"{context_block}\n\n"
                f"Outline to follow:\n{outline}\n\n"
                f"Write the full document following the outline above. "
                f"Complete ALL sections thoroughly.{depth_instruction} {format_instruction}"
            )
            draft = llm_client.complete_text(
                prompt=write_prompt,
                system_message="You are an expert writer. Produce well-structured, comprehensive content.",
                temperature=0.7,
                max_tokens=8192,
                model_override=writer_model,
            )
        else:
            logger.info(
                f"compose_document [{length}]: writing {len(sections)} sections individually"
            )
            section_contents = []
            for section in sections:
                bullets_text = "\n".join(section["bullets"]) if section["bullets"] else ""
                # Provide last 2 written sections as rolling context to maintain coherence
                prior_context = "\n\n".join(section_contents[-2:]) if section_contents else ""
                prior_block = (
                    f"\nPrevious sections (for context and continuity — do NOT repeat them):\n{prior_context}\n"
                    if prior_context
                    else ""
                )
                section_prompt = (
                    f"Topic: {topic}\n"
                    f"Instructions: {instructions}\n"
                    f"{context_block}\n\n"
                    f"Full document outline:\n{outline}\n\n"
                    f"{prior_block}\n"
                    f"Now write ONLY Section {section['number']}: {section['title']}\n"
                    + (f"Key points to cover:\n{bullets_text}\n\n" if bullets_text else "\n")
                    + f"Write ONLY this section's content (do not include other sections).{depth_instruction} {format_instruction}"
                )
                section_content = llm_client.complete_text(
                    prompt=section_prompt,
                    system_message=(
                        "You are an expert writer. Write detailed, high-quality section content. "
                        "Return ONLY the section content — no preamble, no commentary."
                    ),
                    temperature=0.7,
                    max_tokens=section_tokens,
                    model_override=writer_model,
                )
                section_contents.append(section_content)
                logger.info(
                    f"compose_document [{length}]: section {section['number']}/{len(sections)} complete"
                )
            draft = "\n\n".join(section_contents)

    # --- Step 3: Review ---
    logger.info("compose_document: reviewing and improving draft")
    review_prompt = (
        f"Original instructions: {instructions}\n\n"
        f"Draft document:\n{draft}\n\n"
        "Review this draft for clarity, completeness, flow, and correctness. "
        "Return the improved version of the full document. "
        "Do NOT return commentary — only the improved document text."
    )
    final_content = llm_client.complete_text(
        prompt=review_prompt,
        system_message=(
            "You are a meticulous editor. Improve the document while preserving its structure and intent. "
            "Return ONLY the improved document, no meta-commentary."
        ),
        temperature=0.4,
        max_tokens=8192,
        model_override=writer_model,
    )
    logger.info("compose_document: review complete")

    # Count sections from outline
    section_count = len(
        [line for line in outline.split("\n") if re.match(r"^\s*\d+[\.\)]\s", line)]
    )

    # Persist the final composed document as an internal transparency row.
    if conversation_id:
        transparency_logger.log(
            conversation_id, "document", final_content, role="assistant",
            extra_metadata={"stage": "final"},
        )

    return {
        "content": final_content,
        "outline": outline,
        "sections": section_count,
    }


def extract_text_from_bytes(data: bytes, filename: str) -> Dict[str, Any]:
    """
    Extract text content from binary file data based on file extension.

    Supports: PDF (.pdf), Word (.docx), plain text (.txt, .md, .csv, .json, .xml, .html).

    Args:
        data: Raw file bytes
        filename: Original filename (used for extension detection)

    Returns:
        Dict with keys: success (bool), text (str), format (str), message (str)
    """
    ext = os.path.splitext(filename)[1].lower() if filename else ""

    # Plain text formats
    text_extensions = {
        ".txt",
        ".md",
        ".csv",
        ".json",
        ".xml",
        ".html",
        ".htm",
        ".yaml",
        ".yml",
        ".ini",
        ".cfg",
        ".log",
        ".py",
        ".js",
        ".ts",
        ".css",
        ".sql",
        ".sh",
    }
    if ext in text_extensions:
        try:
            text = data.decode("utf-8")
            return {"success": True, "text": text, "format": ext.lstrip("."), "message": ""}
        except UnicodeDecodeError:
            try:
                text = data.decode("latin-1")
                return {"success": True, "text": text, "format": ext.lstrip("."), "message": ""}
            except Exception:
                return {
                    "success": False,
                    "text": "",
                    "format": ext.lstrip("."),
                    "message": "Could not decode text file",
                }

    # PDF - requires Mistral OCR for text extraction (handled by calling code)
    if ext == ".pdf":
        logger.info("PDF extraction requested via document utility - requires Mistral OCR")
        return {
            "success": False,
            "text": "",
            "format": "pdf",
            "message": "PDF text extraction requires Mistral OCR. The calling service should handle OCR.",
        }

    # XLSX (Excel spreadsheets)
    if ext in (".xlsx", ".xls"):
        try:
            from openpyxl import load_workbook

            wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            sheets_text = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    # Skip completely empty rows
                    if any(cells):
                        rows.append("\t".join(cells))
                if rows:
                    sheets_text.append(f"## Sheet: {sheet_name}\n" + "\n".join(rows))
            wb.close()

            text = "\n\n".join(sheets_text) if sheets_text else ""
            if not text:
                return {
                    "success": True,
                    "text": "(empty spreadsheet)",
                    "format": "xlsx",
                    "message": "Spreadsheet contains no data",
                }
            return {
                "success": True,
                "text": text,
                "format": "xlsx",
                "message": f"Extracted text from {len(sheets_text)} sheet(s)",
            }
        except Exception as e:
            logger.error(f"XLSX extraction failed: {e}")
            return {
                "success": False,
                "text": "",
                "format": "xlsx",
                "message": f"XLSX extraction failed: {str(e)}",
            }

    # DOCX
    if ext == ".docx":
        try:
            from docx import Document

            doc = Document(io.BytesIO(data))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            return {
                "success": True,
                "text": text,
                "format": "docx",
                "message": f"Extracted text from {len(paragraphs)} paragraphs",
            }
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return {
                "success": False,
                "text": "",
                "format": "docx",
                "message": f"DOCX extraction failed: {str(e)}",
            }

    # Unknown format
    return {
        "success": False,
        "text": "",
        "format": ext.lstrip(".") or "unknown",
        "message": f"Unsupported file format: {ext or 'unknown'}. "
        f"Supported: pdf, xlsx, docx, txt, md, csv, json, xml, html.",
    }


def create_docx(
    content: str,
    filename: str,
    title: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Create a .docx file from markdown-like content.

    Args:
        content: Markdown-formatted content
        filename: Output filename
        title: Optional document title

    Returns:
        Dict with path to created file and metadata
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    if title:
        doc.add_heading(title, level=0)

    lines = content.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Code block
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$|^\*{3,}$|^_{3,}$", line.strip()):
            doc.add_paragraph("_" * 50)
            i += 1
            continue

        # Bullet list
        if re.match(r"^[-*]\s+", line):
            text = re.sub(r"^[-*]\s+", "", line).strip()
            doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        # Numbered list
        num_match = re.match(r"^\d+[.)]\s+", line)
        if num_match:
            text = line[num_match.end() :].strip()
            doc.add_paragraph(text, style="List Number")
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            quote_text = line[2:].strip()
            p = doc.add_paragraph()
            run = p.add_run(quote_text)
            run.italic = True
            i += 1
            continue

        # Markdown table
        if line.startswith("|") and line.count("|") >= 2:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].count("|") >= 2:
                table_lines.append(lines[i])
                i += 1

            rows = _parse_table_rows(table_lines)
            if rows:
                num_cols = max(len(row) for row in rows)
                table = doc.add_table(rows=len(rows), cols=num_cols)
                table.style = "Table Grid"

                for row_idx, row in enumerate(rows):
                    for col_idx in range(num_cols):
                        cell_text = row[col_idx] if col_idx < len(row) else ""
                        cell = table.cell(row_idx, col_idx)
                        cell.paragraphs[0].clear()
                        _add_formatted_runs(cell.paragraphs[0], cell_text)
                        if row_idx == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True

                doc.add_paragraph()
            continue

        # Regular paragraph - collect consecutive lines
        para_lines = [line.strip()]
        i += 1
        while i < len(lines):
            next_line = lines[i]
            if not next_line.strip():
                break
            if re.match(r"^(#{1,3}\s|[-*]\s|>\s|\d+[.)]\s|```|---+|\*\*\*+|___+)", next_line):
                break
            if next_line.strip().startswith("|") and next_line.count("|") >= 2:
                break
            para_lines.append(next_line.strip())
            i += 1

        full_text = " ".join(para_lines)
        p = doc.add_paragraph()
        _add_formatted_runs(p, full_text)

    # Save to application temp directory
    temp_dir = os.path.join(os.getenv("TMP_DIR", "tmp"), "assistant_docs")
    os.makedirs(temp_dir, exist_ok=True)

    if not filename.endswith(".docx"):
        filename = f"{filename}.docx"

    filepath = os.path.join(temp_dir, filename)
    doc.save(filepath)

    logger.info(f"Created DOCX: {filepath}")

    return {
        "filepath": filepath,
        "filename": filename,
        "size": os.path.getsize(filepath),
        "message": f"Document '{filename}' created successfully ({os.path.getsize(filepath)} bytes).",
    }


def create_pdf(
    content: str,
    filename: str,
    title: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Create a PDF file from markdown content.

    Converts markdown to HTML using the markdown library, then renders it
    to PDF using fpdf2's HTML renderer.

    Args:
        content: Markdown-formatted content
        filename: Output filename
        title: Optional document title

    Returns:
        Dict with path to created file and metadata
    """
    import markdown
    from fpdf import FPDF

    # Prepend title as an H1 if provided
    if title:
        content = f"# {title}\n\n{content}"

    # Convert markdown to HTML
    md = markdown.Markdown(extensions=["tables", "fenced_code"])
    html_body = md.convert(content)

    # Wrap in minimal HTML document
    html = f"""<!DOCTYPE html>
<html>
<body>
{html_body}
</body>
</html>"""

    # Build PDF with Unicode font support.
    # DejaVu Sans has broad Unicode coverage; Liberation Sans is a fallback.
    # Both are installed in the container via Dockerfile apt packages.
    _FONT_CANDIDATES = [
        # (name, dir, regular, bold, italic, bold_italic)
        (
            "DejaVu Sans",
            "/usr/share/fonts/truetype/dejavu",
            "DejaVuSans.ttf",
            "DejaVuSans-Bold.ttf",
            "DejaVuSans-Oblique.ttf",
            "DejaVuSans-BoldOblique.ttf",
        ),
        (
            "Liberation Sans",
            "/usr/share/fonts/truetype/liberation",
            "LiberationSans-Regular.ttf",
            "LiberationSans-Bold.ttf",
            "LiberationSans-Italic.ttf",
            "LiberationSans-BoldItalic.ttf",
        ),
    ]

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    for _name, _dir, _reg, _bold, _ital, _bi in _FONT_CANDIDATES:
        if all(os.path.isfile(os.path.join(_dir, f)) for f in (_reg, _bold, _ital, _bi)):
            pdf.add_font(_name, "", os.path.join(_dir, _reg), uni=True)
            pdf.add_font(_name, "B", os.path.join(_dir, _bold), uni=True)
            pdf.add_font(_name, "I", os.path.join(_dir, _ital), uni=True)
            pdf.add_font(_name, "BI", os.path.join(_dir, _bi), uni=True)
            pdf.set_font(_name, size=12)
            break
    else:
        pdf.set_font("Helvetica", size=12)

    # Register a Unicode-capable monospace font so that <pre>/<code> blocks
    # (e.g. ASCII diagrams with box-drawing characters) don't fall back to the
    # built-in "courier" font, which has no Unicode coverage.
    _MONO_CANDIDATES = [
        (
            "DejaVuSansMono",
            "/usr/share/fonts/truetype/dejavu",
            "DejaVuSansMono.ttf",
            "DejaVuSansMono-Bold.ttf",
            "DejaVuSansMono-Oblique.ttf",
            "DejaVuSansMono-BoldOblique.ttf",
        ),
    ]
    _mono_font_name = None
    for _name, _dir, _reg, _bold, _ital, _bi in _MONO_CANDIDATES:
        if all(os.path.isfile(os.path.join(_dir, f)) for f in (_reg, _bold, _ital, _bi)):
            pdf.add_font(_name, "", os.path.join(_dir, _reg), uni=True)
            pdf.add_font(_name, "B", os.path.join(_dir, _bold), uni=True)
            pdf.add_font(_name, "I", os.path.join(_dir, _ital), uni=True)
            pdf.add_font(_name, "BI", os.path.join(_dir, _bi), uni=True)
            _mono_font_name = _name
            break

    _write_html_kwargs = {}
    if _mono_font_name:
        from fpdf.html import FontFace

        _code_face = FontFace(family=_mono_font_name)
        _write_html_kwargs["tag_styles"] = {"pre": _code_face, "code": _code_face}

    pdf.write_html(html, **_write_html_kwargs)

    # Save to application temp directory
    temp_dir = os.path.join(os.getenv("TMP_DIR", "tmp"), "assistant_docs")
    os.makedirs(temp_dir, exist_ok=True)

    if not filename.endswith(".pdf"):
        filename = f"{filename}.pdf"

    filepath = os.path.join(temp_dir, filename)
    pdf.output(filepath)

    logger.info(f"Created PDF: {filepath}")

    return {
        "filepath": filepath,
        "filename": filename,
        "size": os.path.getsize(filepath),
        "message": f"PDF '{filename}' created successfully ({os.path.getsize(filepath)} bytes). "
        f"Use nextcloud_upload_file (with source_path) or onedrive_upload_file (with source_path) to upload to cloud storage.",
    }


async def create_html(
    prompt: str,
    filename: str,
    context: Optional[str] = None,
    settings_service=None,
) -> Dict[str, Any]:
    """
    Generate a complete HTML page from a plain-English description.

    Calls the LLM directly to write the HTML. Page-specific CSS and JavaScript
    are inline; CDN chart libraries (Chart.js, Plotly.js, etc.) are permitted
    when charts are needed. The result is saved to temporary storage.

    Args:
        prompt: Plain-English description of the page to generate
        filename: Output filename (will have .html appended if missing)
        context: Optional data or content to embed (CSV, JSON, text, etc.)
        settings_service: SettingsService instance for LLM configuration

    Returns:
        Dict with filepath, filename, size, message
    """
    if not settings_service:
        raise ValueError("settings_service is required for create_html")

    from src.core.llm_client import LLMClient, LLMConfig

    llm_config = LLMConfig.from_settings(settings_service)
    llm_client = LLMClient(llm_config)

    system_message = (
        "You are an expert front-end developer. Generate a complete HTML page.\n"
        "Rules:\n"
        "- All page-specific CSS must be inside a <style> tag in <head>.\n"
        "- All page-specific JavaScript must be inside a <script> tag.\n"
        "- You MAY load chart/visualisation libraries from a CDN when charts are needed "
        "(e.g. Chart.js from cdnjs, Plotly.js from cdn.plot.ly). "
        "Do NOT link to arbitrary external stylesheets or scripts beyond these.\n"
        "- When using a CDN library, embed all data as inline JS variables — "
        "never fetch data from external URLs.\n"
        "- Output ONLY raw HTML starting with <!DOCTYPE html>. "
        "No markdown, no explanation, no code fences.\n"
        "- Make it visually polished: clean typography, good spacing, modern responsive layout.\n"
        "- If data is provided, embed it directly "
        "(e.g. as a JS variable passed to a chart, or rendered as an HTML table/list)."
    )

    user_prompt = f"Build this HTML page: {prompt.strip()}"
    if context:
        user_prompt += f"\n\nData / content to embed:\n{context.strip()}"

    _MAX_TOKENS_PER_CALL = 16000
    _MAX_CONTINUATIONS = 4

    messages = [
        {"role": "system", "content": system_message},
        {"role": "user", "content": user_prompt},
    ]

    html_parts: list[str] = []
    for _ in range(_MAX_CONTINUATIONS):
        response = llm_client.complete(messages=messages, max_tokens=_MAX_TOKENS_PER_CALL)
        chunk = response.choices[0].message.content or ""
        finish_reason = response.choices[0].finish_reason

        # Strip code fences only on the first chunk
        if not html_parts:
            chunk = chunk.strip()
            if chunk.startswith("```"):
                chunk = re.sub(r"^```[a-zA-Z]*\n?", "", chunk)
                chunk = re.sub(r"\n?```$", "", chunk)
                chunk = chunk.strip()

        html_parts.append(chunk)

        if finish_reason != "length":
            break

        # Truncated — ask the model to continue from where it stopped
        messages.append({"role": "assistant", "content": chunk})
        messages.append({"role": "user", "content": "Continue exactly from where you left off."})

    html = "".join(html_parts)

    # Final fence cleanup in case the last chunk ended with one
    html = re.sub(r"\n?```\s*$", "", html).strip()

    from src.utils.tmp import get_tmp_dir

    temp_dir = os.path.join(str(get_tmp_dir()), "assistant_docs")
    os.makedirs(temp_dir, exist_ok=True)

    if not filename.endswith(".html"):
        filename = f"{filename}.html"

    filepath = os.path.join(temp_dir, filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(html)

    size = os.path.getsize(filepath)
    logger.info(f"Created HTML: {filepath}")

    return {
        "filepath": filepath,
        "filename": filename,
        "size": size,
        "message": f"HTML page '{filename}' created successfully ({size} bytes). "
        "Use nextcloud_upload_file (with source_path) or onedrive_upload_file (with source_path) "
        "to upload to cloud storage.",
    }


def _add_formatted_runs(paragraph, text: str):
    """Add text with basic inline formatting (bold, italic) to a paragraph."""
    import re as _re

    pattern = _re.compile(
        r"(\*\*(.+?)\*\*)"  # bold
        r"|(\*(.+?)\*)"  # italic
    )

    last_end = 0
    for match in pattern.finditer(text):
        if match.start() > last_end:
            paragraph.add_run(text[last_end : match.start()])

        if match.group(2):  # bold
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(4):  # italic
            run = paragraph.add_run(match.group(4))
            run.italic = True

        last_end = match.end()

    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def _parse_table_rows(table_lines: list) -> list:
    """Parse raw markdown table lines into a list of rows (list of cell strings). Separator rows are skipped."""
    rows = []
    for line in table_lines:
        stripped = line.strip().strip("|")
        cells = [cell.strip() for cell in stripped.split("|")]
        if all(re.match(r"^[-: ]+$", cell) for cell in cells if cell):
            continue
        rows.append(cells)
    return rows
